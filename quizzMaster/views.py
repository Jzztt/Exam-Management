from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework_simplejwt.authentication import JWTAuthentication
from rest_framework.views import APIView
from django.contrib.auth.hashers import make_password
from rest_framework_simplejwt.tokens import RefreshToken
from .models import Choice, Exam, Question, Subject, ExamSchedule,CorrectAnswer, ExamQuestion, User, Submission, Role, UserSubject
from .serializers import ChoiceSerializer, ExamSerializer, QuestionSerializer, SubjectSerializer, ExamScheduleSerializer, ExamQuestionSerializer, UserSerializer, SubmissionSerializer, RoleSerializer, UserSubjectSerializer
from .permissions import IsStudent, IsAdmin, IsExamAdministrator, IsQuestionManager
from rest_framework.exceptions import ValidationError
from docx import Document
import cloudinary
from cloudinary.uploader import upload


class ProtectedView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        return Response({"message": "Protected data"}, status=status.HTTP_200_OK)
class ChoiceViewSet(viewsets.ModelViewSet):
    queryset = Choice.objects.all()
    serializer_class = ChoiceSerializer
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsStudent]

class QuestionViewSet(viewsets.ModelViewSet):
    queryset = Question.objects.all()
    serializer_class = QuestionSerializer
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsQuestionManager | IsAdmin]

class ExamViewSet(viewsets.ModelViewSet):
    queryset = Exam.objects.all()
    serializer_class = ExamSerializer
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsAdmin | IsExamAdministrator]


class ExamQuestionViewSet(viewsets.ModelViewSet):
    queryset = ExamQuestion.objects.all()
    serializer_class = ExamQuestionSerializer
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsAdmin | IsExamAdministrator | IsQuestionManager]

class SubjectViewSet(viewsets.ModelViewSet):
    queryset = Subject.objects.all()
    serializer_class = SubjectSerializer
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated | IsAdmin | IsQuestionManager]

    def create(self, request, *args, **kwargs):
        serializer = SubjectSerializer(data=request.data)
        try:
            serializer.is_valid(raise_exception=True)
            self.perform_create(serializer)
            return Response(
                {"message": "Subject created successfully!", "data": serializer.data},
                status=status.HTTP_201_CREATED
            )
        except ValidationError as e:
            error_details = e.detail
            return Response(
                {"error": str(error_details), "message": "Validation error occurred."},
                status=status.HTTP_400_BAD_REQUEST
            )

class ExamScheduleViewSet(viewsets.ModelViewSet):
    queryset = ExamSchedule.objects.all()
    serializer_class = ExamScheduleSerializer
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsAdmin | IsExamAdministrator]

class SubmissionViewSet(viewsets.ModelViewSet):
    queryset = Submission.objects.all()
    serializer_class = SubmissionSerializer
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsStudent]

class RoleViewSet(viewsets.ModelViewSet):
    queryset = Role.objects.all()
    serializer_class = RoleSerializer
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsAdmin]

class UserSubjectViewSet(viewsets.ModelViewSet):
    queryset = UserSubject.objects.all()
    serializer_class = UserSubjectSerializer
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsAdmin]

class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsAdmin]

class ImportDocxView(APIView):
    # authentication_classes = [JWTAuthentication]
    # permission_classes = [IsAuthenticated, IsAdmin, IsQuestionManager]
    def save_to_database( self, questions, exam_id, subject_id):
        for question in questions:
            question_text = question['question']
            choices = question['choices']
            correct_choice = question.get('correct_choice')

            subject = Subject.objects.get(id=subject_id)
            question = Question.objects.create(question_text=question_text, subject=subject)
            correct_choice_obj = None
            for choice_data in choices:
                choice_text = choice_data["text"]
                choice = Choice.objects.create(question=question, choice_text=choice_text)
                if choice_data["option"] == correct_choice:
                    correct_choice_obj = choice
            if correct_choice_obj:
                CorrectAnswer.objects.create(question=question, choice=correct_choice_obj)
            exam = Exam.objects.get(id=exam_id)
            ExamQuestion.objects.create(exam=exam, question=question)

    def post(self, request, *args, **kwargs):
        exam_id = request.data.get('exam_id')
        subject_id = request.data.get('subject_id')
        docx_file = request.FILES.get('file')
        questions = []
        current_question = None
        current_choices = []
        correct_choice = None
        if not docx_file or not docx_file.name.endswith('.docx'):
            return Response({"message": "Invalid file format", "status": "error"}, status=400)
        if not exam_id or not subject_id:
            return Response({"message": "Missing exam_id or subject_id", "status": "error"}, status=400)
        try:
            document = Document(docx_file)
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text.startswith("Câu"):
                    if current_question and current_choices:
                        questions.append({
                            "question": current_question,
                            "choices": current_choices
                        })
                    question_text = text.split(":")[1].strip()
                    current_question = question_text
                    current_choices = []
                elif text.startswith(("A.", "B.", "C.", "D.")):
                    option,choice_text = text.split(".", 1)
                    current_choices.append({
                        "option": option.strip(),
                        "text": choice_text.strip()
                    })
                elif text.startswith("Đáp án"):
                    correct_choice = text.split(":")[1].strip()
                    if current_choices and current_question:
                        questions.append({
                            "question": current_question,
                            "choices": current_choices,
                            "correct_choice": correct_choice
                        })
                        current_question = None
                        current_choices = []
                        correct_choice = None
            if current_question and current_choices:
                questions.append({
                    "question": current_question,
                    "choices": current_choices
                })
            self.save_to_database(questions, exam_id, subject_id)
        except Exception as e:
            return Response({"message": "Something went wrong", "status": "error"}, status=400)
        else:
            return Response({"message": "File imported successfully", "status": "success"}, status=200)
import datetime
class ImportExamView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated, IsAdmin | IsQuestionManager | IsExamAdministrator]

    def extract_metadata(self,document):
        subject_name = None
        num_questions = 0
        lecturer = None

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text.startswith("Subject:"):
                subject_name = text.split(":")[1].strip()
            elif text.startswith("Number of Quiz:"):
                num_questions = int(text.split(":")[1].strip())
            elif text.startswith("Lecturer:"):
                lecturer = text.split(":")[1].strip()

        return subject_name, num_questions, lecturer
    def create_or_update_subject(self,subject_name, lecturer):
        subject, created = Subject.objects.get_or_create(name=subject_name)
        if created:
            subject.lecturer = lecturer
            subject.save()
        return subject


    def upload_image_from_docx(self ,doc, image_reference):
        from docx import Document
        from cloudinary.uploader import upload
        from cloudinary.utils import cloudinary_url
        from io import BytesIO
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image = rel.target_part.blob
                    result = upload(BytesIO(image), resource_type="image",upload_preset='quiz_images_upload')
                    print(result["secure_url"])
                    return result["secure_url"]
        except Exception as e:
            print(f"Error uploading image {image_reference}: {e}")
            return None

    def process_questions_table(self,table, doc):

        questions = []
        current_question = None
        correct_choice = None
        mark = None
        unit = None
        mix_choices = False
        for row in table.rows:
            first_cell = row.cells[0].text.strip()
            second_cell = row.cells[1].text.strip()
            if first_cell.startswith("QN"):
                image_file = None

                if "[file:" in second_cell:
                    question_text = second_cell.split("[file:")[0].strip()
                    image_reference  = second_cell.split("[file:")[1].split("]")[0].strip()
                    image_file = self.upload_image_from_docx(doc, image_reference)
                else:
                    question_text = second_cell

                current_question = {
                    "question_text": question_text,
                    "choices": [],
                    "image_file": image_file,
                    "correct_choice": correct_choice,
                    "mark": mark,
                    "unit": unit,
                    "mix_choices": mix_choices
                }

                questions.append(current_question)

            elif first_cell.startswith(("a.", "b.", "c.", "d.")):
                option = first_cell.split(".", 1)[0]
                choice_text = second_cell
                if current_question:
                     current_question["choices"].append({
                        "option": option,
                        "text": choice_text
                    })
            elif first_cell.startswith("ANSWER:"):
                correct_choice = second_cell.strip()
                if current_question:
                    current_question["correct_choice"] = correct_choice
            elif first_cell.startswith("MARK:"):
                mark = float(second_cell.strip())
                if current_question:
                    current_question["mark"] = mark
            elif first_cell.startswith("UNIT:"):
                unit = second_cell.strip()
                if current_question:
                    current_question["unit"] = unit
            elif first_cell.startswith("MIX CHOICES:"):
                mix_choices = second_cell.strip().lower() == "yes"
                if current_question:
                    current_question["mix_choices"] = mix_choices
        return questions

    def save_exam(self,subject, num_questions, exam_code, exam_id = None):
        if exam_id:
            print("exam_id",exam_id)
            exam = Exam.objects.get(id=exam_id)
            print("exam",exam)
            exam.subject = subject
            exam.num_questions = num_questions
            exam.exam_code = exam_code
            exam.save()
            return exam
        else:
            exam = Exam.objects.create(subject=subject, num_questions=num_questions, exam_code=exam_code)
            return exam

    def save_questions_and_choices_and_answers(self, questions, exam, subject):
        for question_data in questions:
            question_text = question_data["question_text"]
            choices = question_data["choices"]
            image_file = question_data.get("image_file")
            unit = question_data.get("unit")
            is_mixed = question_data.get("mix_choices")
            correct_answer = question_data.get("correct_choice")

            question = Question.objects.create(question_text=question_text, image=image_file, unit=unit, is_mixed=is_mixed, subject=subject)
            for choice_data in choices:
                option = choice_data["option"]
                choice_text = choice_data["text"]
                Choice.objects.create(question=question, choice_text=choice_text, option=option)
            if correct_answer:
                correct_choice = Choice.objects.get(question=question, option=correct_answer)
                CorrectAnswer.objects.create(question=question, choice=correct_choice)
            ExamQuestion.objects.create(exam=exam, question=question)

    def post(self, request, *args, **kwargs):
        docx_file = request.FILES.get('file')

        if not docx_file or not docx_file.name.endswith('.docx'):
            return Response({"message": "Invalid file format", "status": "error"}, status=400)
        try:
            document = Document(docx_file)

            subject_name, num_questions, lecturer = self.extract_metadata(document)

            if not subject_name or not num_questions or not lecturer:
                return Response({"message": "Missing required fields", "status": "error"}, status=400)

            subject = self.create_or_update_subject(subject_name, lecturer)

            exam_id  = request.data.get('exam_id')
            if exam_id:
                try:
                    exam = Exam.objects.get(id=exam_id)
                    exam_code = exam.exam_code
                except Exam.DoesNotExist:
                    return Response({"message": "Exam not found", "status": "error"}, status=404)
            else:
                exam_code = f"EXAM_{subject_name}_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"

            exam = self.save_exam(subject, num_questions, exam_code, exam_id)

            table = document.tables[0]
            questions = self.process_questions_table(table,document)
            print(questions)

            self.save_questions_and_choices_and_answers(questions, exam, subject)

        except Exception as e:
            return Response({"message": str(e), "status": "error"}, status=400)

        return Response({"message": "File imported successfully", "status": "success"}, status=200)
class AuthViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer

    @action(detail=False, methods=['post'], permission_classes=[AllowAny])
    def register(self, request):
       username = request.data.get('username')
       email = request.data.get('email')
       password = request.data.get('password')

       try:
        role = Role.objects.get(name='Candidate')
       except Role.DoesNotExist:
           return Response({'message': 'Role not found', "status": "error"}, status=status.HTTP_400_BAD_REQUEST)

       if User.objects.filter(username=username).exists():
            return Response({'message': 'username already exists', 'status': 'error'}, status=status.HTTP_400_BAD_REQUEST)
       if User.objects.filter(email=email).exists():
            return Response({'message': 'Email already exists', 'status': 'error'}, status=status.HTTP_400_BAD_REQUEST)

       user = User.objects.create(username=username, email=email, password=make_password(password), role=role)
       user.save()
       return Response({'message': 'User registered successfully', 'data': UserSerializer(user).data, 'status': 'success'}, status=status.HTTP_201_CREATED)

    @action(detail=False, methods=['post'], permission_classes=[AllowAny])
    def login(self, request):
        username = request.data.get('username')
        password = request.data.get('password')
        try:
            user = User.objects.get(username=username)
            if user.check_password(password):
                refresh = RefreshToken.for_user(user)
                return Response({
                    'refresh': str(refresh),
                    'access': str(refresh.access_token),
                    'username': {
                        'id': user.id,
                        'username': user.username,
                        'email': user.email
                    },
                    'message': 'User logged in successfully',
                    'status': 'success'
                }, status=status.HTTP_200_OK)
            else:
                return Response({'message': 'Invalid credentials', 'status': 'error'}, status=status.HTTP_400_BAD_REQUEST)
        except User.DoesNotExist:
            return Response({'message': 'User does not exist', 'status': 'error'}, status=status.HTTP_404_NOT_FOUND)

    @action(detail=False, methods=['get'], permission_classes=[IsAuthenticated])
    def me(self, request):
        user = request.user
        serializer = UserSerializer(user)
        return Response({'user': serializer.data, 'message': 'User details', 'status': 'success'}, status=status.HTTP_200_OK)


